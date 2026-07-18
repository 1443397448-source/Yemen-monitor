"""Markdown → DOCX 转换，处理标题/粗体/链接/列表等常见元素。"""
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

# 需按顺序匹配的行模式
HEADING = re.compile(r"^(#{1,4})\s+(.+)$")
HR = re.compile(r"^──+$|^---$|^\*\*\*$")
BULLET = re.compile(r"^-\s+(.+)$")
LINK = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
BOLD = re.compile(r"\*\*(.+?)\*\*")

FONT_NAME = "微软雅黑"
FONT_SIZE = 10.5


def md_to_docx(md_path, docx_path):
    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(FONT_SIZE)
    font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)

    with open(md_path, encoding="utf-8") as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # 空行
        if not line:
            i += 1
            continue

        # 水平线/分隔符
        if HR.match(line):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # 标题
        hm = HEADING.match(line)
        if hm:
            level = min(len(hm.group(1)), 4)
            text = hm.group(2)
            h = doc.add_heading(level=level)
            _add_rich_text(h, text)
            i += 1
            continue

        # 无序列表
        bm = BULLET.match(line)
        if bm:
            p = doc.add_paragraph(style="List Bullet")
            _add_rich_text(p, bm.group(1))
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        _add_rich_text(p, line)
        i += 1

    doc.save(docx_path)
    return docx_path


def _add_rich_text(paragraph, text):
    """把 **粗体** 和 [文字](链接) 混排到段落中。"""
    parts = []
    # 先拆分 link，保护 URL 不被 bold 正则误匹配
    pos = 0
    for lm in LINK.finditer(text):
        if lm.start() > pos:
            parts.append(("text", text[pos:lm.start()]))
        parts.append(("link", lm.group(1), lm.group(2)))
        pos = lm.end()
    if pos < len(text):
        parts.append(("text", text[pos:]))

    for part in parts:
        if part[0] == "text":
            _bold_segments(paragraph, part[1])
        elif part[0] == "link":
            _add_hyperlink(paragraph, part[1], part[2])


def _bold_segments(paragraph, text):
    pos = 0
    for bm in BOLD.finditer(text):
        if bm.start() > pos:
            paragraph.add_run(text[pos:bm.start()])
        run = paragraph.add_run(bm.group(1))
        run.bold = True
        pos = bm.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _add_hyperlink(paragraph, text, url):
    """在 DOCX 中添加可点击的超链接（需操作底层 XML）。"""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                          is_external=True)
    from docx.oxml import OxmlElement
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run_el = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), "0563C1")
    rpr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(u)
    run_el.append(rpr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run_el.append(t)
    hyperlink.append(run_el)
    paragraph._p.append(hyperlink)
